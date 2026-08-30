"""解析 SRM 蓝图 docx：保留正文段落与表格（整表转为多行文本），并记录出现顺序。

返回 elements：list of ("p", text) 或 ("table", [row_str, ...])。
"""
from __future__ import annotations

from pathlib import Path
from typing import List, Tuple

from docx import Document

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _para_text(p_el) -> str:
    return "".join(t.text or "" for t in p_el.iter(f"{{{W}}}t"))


def parse_docx(path: str) -> List[Tuple[str, object]]:
    doc = Document(str(path))
    body = doc.element.body
    elements: List[Tuple[str, object]] = []
    tbl_index = 0
    for child in body.iterchildren():
        tag = child.tag.split("}")[-1]
        if tag == "p":
            text = _para_text(child).strip()
            if text:
                elements.append(("p", text))
        elif tag == "tbl":
            if tbl_index < len(doc.tables):
                table = doc.tables[tbl_index]
                rows = []
                for r in table.rows:
                    cells = [("".join(c.text or "" for c in cell.paragraphs)) for cell in r.cells]
                    rows.append(" | ".join(cells))
                tbl_index += 1
                if rows:
                    elements.append(("table", rows))
    return elements
